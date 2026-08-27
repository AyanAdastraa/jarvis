import os
import csv
import logging
from typing import Optional
from pathlib import Path
from pypdf import PdfReader
import docx

logger = logging.getLogger(__name__)

class DocumentParseError(Exception):
    pass

class ParsedDocument:
    def __init__(self, filename: str, mime_type: str, content: str, metadata: dict = None):
        self.filename = filename
        self.mime_type = mime_type
        self.content = content
        self.metadata = metadata or {}

def _parse_pdf(file_path: Path) -> str:
    text = []
    try:
        reader = PdfReader(str(file_path))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n\n".join(text)
    except Exception as e:
        raise DocumentParseError(f"Failed to parse PDF: {e}")

def _parse_docx(file_path: Path) -> str:
    try:
        doc = docx.Document(str(file_path))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise DocumentParseError(f"Failed to parse DOCX: {e}")

def _parse_csv(file_path: Path) -> str:
    try:
        with open(file_path, mode='r', encoding='utf-8') as f:
            reader = csv.reader(f)
            return "\n".join([", ".join(row) for row in reader])
    except Exception as e:
        raise DocumentParseError(f"Failed to parse CSV: {e}")

def _parse_text(file_path: Path) -> str:
    try:
        with open(file_path, mode='r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise DocumentParseError(f"Failed to parse text file: {e}")

def parse_document(file_path: str, mime_type: Optional[str] = None) -> ParsedDocument:
    path = Path(file_path)
    if not path.exists():
        raise DocumentParseError(f"File not found: {file_path}")
        
    ext = path.suffix.lower()
    
    # Simple extension to mime type mapping if not provided
    if not mime_type:
        if ext == '.pdf':
            mime_type = 'application/pdf'
        elif ext == '.docx':
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif ext == '.csv':
            mime_type = 'text/csv'
        elif ext == '.md':
            mime_type = 'text/markdown'
        elif ext == '.txt':
            mime_type = 'text/plain'
        else:
            mime_type = 'application/octet-stream'

    content = ""
    if ext == '.pdf':
        content = _parse_pdf(path)
    elif ext == '.docx':
        content = _parse_docx(path)
    elif ext == '.csv':
        content = _parse_csv(path)
    elif ext in ['.txt', '.md']:
        content = _parse_text(path)
    else:
        raise DocumentParseError(f"Unsupported file type/extension: {ext}")
        
    return ParsedDocument(
        filename=path.name,
        mime_type=mime_type,
        content=content.strip(),
        metadata={"extension": ext}
    )
